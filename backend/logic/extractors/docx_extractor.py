# import docx

# def extract_text_from_docx(file_path):
#     """Extract full text from DOCX"""
#     doc = docx.Document(file_path)
#     texts = [p.text for p in doc.paragraphs]
#     return [{"page_number": 1, "full_text": "\n".join(texts), "file_path": file_path}]


# backend/logic/docx_extractor.py
import os
import uuid
import logging
import tempfile
from pathlib import Path
from typing import List, Dict, Any
 
import docx
from .image_extractor import extract_text_from_image
 
logger = logging.getLogger(__name__)
 
def _safe_basename(path):
    try:
        return os.path.basename(path)
    except Exception:
        return str(path)
 
def extract_text_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Robust DOCX extractor.
    Returns a list of dicts:
      {"file_name", "file_path", "page_number", "full_text", "status"}
    Additionally returns any embedded image extractions as separate entries.
    """
    results = []
    file_name = _safe_basename(file_path)
    p = Path(file_path)
 
    if not p.exists():
        return [{"file_name": file_name, "file_path": file_path, "page_number": 0, "full_text": "", "status": "error", "error": "file_not_found"}]
 
    try:
        doc = docx.Document(str(p))
    except Exception as e:
        logger.exception("failed to open docx: %s", e)
        return [{"file_name": file_name, "file_path": file_path, "page_number": 0, "full_text": "", "status": "error", "error": str(e)}]
 
    try:
        # Paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    paragraphs.append(row_text)
 
        combined_text = "\n".join(paragraphs).strip()
 
        # Add textual result (DOCX doesn't have reliable "pages" — keep page_number 1)
        results.append({
            "file_name": file_name,
            "file_path": file_path,
            "page_number": 1,
            "full_text": combined_text,
            "status": "ok"
        })
 
        # If the doc has embedded images, extract text from them as fallback (use only when helpful)
        # Many images live in rels with reltype containing "image"
        try:
            # iterate relationships for images
            rels = getattr(doc.part, "rels", {}) or {}
            for rel in list(rels.values()):
                try:
                    if getattr(rel, "reltype", "").lower().endswith("/image") or "image" in str(getattr(rel, "reltype", "")).lower():
                        part = getattr(rel, "target_part", None)
                        if not part:
                            continue
                        blob = getattr(part, "blob", None)
                        if not blob:
                            continue
                        # write blob to temp file
                        ext = ""
                        # try to infer extension
                        content_type = getattr(part, "content_type", "") or ""
                        if "/" in content_type:
                            ext = "." + content_type.split("/")[-1]
                        tmp_name = f"docx_emb_{uuid.uuid4().hex}{ext or '.png'}"
                        tmp_path = Path(tempfile.gettempdir()) / tmp_name
                        with open(tmp_path, "wb") as f:
                            f.write(blob)
                        # call image extractor (returns list)
                        try:
                            img_entries = extract_text_from_image(str(tmp_path))
                            # append each entry to results
                            for e in img_entries:
                                # ensure file_name and file_path are present
                                e.setdefault("file_name", tmp_path.name)
                                e.setdefault("file_path", str(tmp_path))
                                results.append(e)
                        except Exception as ie:
                            logger.exception("image extraction from embedded failed: %s", ie)
                        finally:
                            try:
                                tmp_path.unlink()
                            except Exception:
                                pass
                except Exception:
                    # ignore single rel failures
                    continue
        except Exception as e:
            logger.exception("embedded image extraction loop failed: %s", e)
 
        return results
 
    except Exception as e:
        logger.exception("docx extraction unexpected error: %s", e)
        return [{"file_name": file_name, "file_path": file_path, "page_number": 1, "full_text": "", "status": "error", "error": str(e)}]
 
 